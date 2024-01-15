
import csv
from email import message
from http import client
from lib2to3.fixes.fix_except import find_excepts
from math import fabs
from pickle import TRUE



def read_specific_cell(file_path, row_index, column_index):
    with open(file_path, newline="", encoding="UTF-8") as csvfile:
        csv_reader = csv.reader(csvfile)

        # 使用 next 方法跳到指定行
        for _ in range(row_index):
            next(csv_reader)

        # 获取指定行的数据
        row = next(csv_reader, None)

        if row is not None and column_index < len(row):
            # 获取指定列的数据
            cell_value = row[column_index]
            # print(f"在第 {row_index+1} 行、第 {column_index+1} 列的单元格中的信息是：{cell_value}")
            # print(f"get row : {row_index} column : {column_index}")
            return row[column_index]
        else:
            print("指定的行或列超出范围")


def get_row_number(file_path):
    with open(file_path, newline="", encoding="UTF-8") as csvfile:
        csv_reader = csv.reader(csvfile)
        # 获取行数和列数
        row_count = sum(1 for row in csv_reader)
        return row_count
    
def compare_prefix(str1, str2, prefix_length):
    return str1[:prefix_length] == str2[:prefix_length]


import time
import paho.mqtt.client as mqtt
client = mqtt.Client()
# 回调函数 - 当客户端连接到服务器时调用
def on_connect(client, userdata, flags, rc):
    if rc == 0:
        print("mqtt连接成功")
        # client.subscribe("/nhh/python")
    else:
        print(f"连接失败，返回码： {rc}")


mqtt_msg = "test_msg"
flag = TRUE
DeviceComponentList = "test_msg"

# 回调函数 - 当接收到消息时调用
def on_message(client, userdata, msg):
    topic = msg.topic
    message = msg.payload.decode('utf-8',errors = "ignore")
    # if message =="T01NOnline" or message == "T01NOffline" or message == "T01NMessageFormartError" or message[:3] == "S01M" :
    # if message[:4] == "J01N" or message[:4] == "T01N":
    if message[:4] == "J01N" :
        json_msg =json.loads(message[4:])
        if json_msg["_type"] == "DeviceComponentList" :
            global DeviceComponentList
            DeviceComponentList = message
        else:
            print("get recv\n")
            global mqtt_msg 
            global flag
            flag = not flag
            mqtt_msg = message
    else :
        print(f"接收到消息'{topic}': {message[:4]}")
       

def mqtt_subscribe(topic):
    client.subscribe(topic)
    print(f"订阅：{topic}\n")

def mqtt_publish(topic,msg):
    client.publish(topic,msg)

def mqtt_init(broker, port, user, paswd):
    client.on_connect = on_connect
    client.on_message = on_message
    broker_address = broker
    broker_port = 1883
    username = user
    password = paswd
    client.username_pw_set(username, password)
    client.connect(broker_address, broker_port, keepalive=60)
    client.loop_start()
    # time.sleep(1)

def mqtt_deinit() :
    client.disconnect()
    print("mqtt连接断开")

def get_flag():
    global flag
    return flag
    
def get_mqtt_msg():
    global mqtt_msg
    return   mqtt_msg

def get_DeviceComponentList():
    global DeviceComponentList
    return DeviceComponentList 

import os
import json

def get_info_from_lic(f_path):
    folder_path = f_path  # 替换为你的文件夹路径
    file_type = ".lic"  # 替换为你的文件类型，例如 ".txt", ".csv", 等

    # 获取文件夹中所有文件
    all_files = os.listdir(folder_path)

    # 筛选出指定类型的文件
    filtered_files = [file for file in all_files if file.endswith(file_type)]

    # 检查是否有符合条件的文件
    if len(filtered_files) == 1:
        # 如果只有一个文件符合条件，打开它
        file_path = os.path.join(folder_path, filtered_files[0])
        try:
            with open(file_path, 'r') as file:
                content = file.read()
                json_data = json.loads(content)
                device_id = json_data["messageId"]
                device_model = json_data["modelInfo"]["model"]
                sn = json_data["sn"]
                return device_id,device_model,sn
        except Exception as e:
            print(f"An error occurred while opening the file: {str(e)}")
    elif len(filtered_files) == 0:
        print(f"No {file_type} files found in the folder.")
    else:
        print(f"Multiple {file_type} files found in the folder. Specify a unique file.")


def set_slot(task_cmd,Otdr_slot,OpticSwitcher_slot,Rcu_slot,PowerMeter_slot) :
    cmd_temp = task_cmd
    json_cmd = json.loads(cmd_temp[4:])
    if  json_cmd["_type"] == "StartOtdrTestAndSwitchSwitcher" :
        json_cmd["target"]["slot"] = Otdr_slot
        json_cmd["switcherConfigs"][0]["slot"] = OpticSwitcher_slot
        cmd_temp = "J01N" + json.dumps(json_cmd)
        return cmd_temp
    elif json_cmd["_type"] == "SwitchSwitcher" :
        json_cmd["target"]["slot"] = OpticSwitcher_slot
        cmd_temp = "J01N" + json.dumps(json_cmd)
        return cmd_temp
    elif json_cmd["_type"] == "StopReportOpticalPower" :
        json_cmd["target"]["Slot"] = PowerMeter_slot
        cmd_temp = "J01N" + json.dumps(json_cmd)
        return cmd_temp
    return cmd_temp

from datetime import date, datetime, timedelta

def get_nowtimestamp_cmd(cmd) :
    cmd_temp = cmd
    json_cmd = json.loads(cmd_temp[4:])
    if json_cmd["_type"] == "SetTime" :
    # 获取当前时间戳
        timestamp = datetime.timestamp(datetime.now())
        json_cmd["timestamp"] = timestamp
        cmd_temp = "J01N" + json.dumps(json_cmd)
        return cmd_temp
    else:
        return cmd_temp


def check_for_ERROR(file_path,error_str):
    try:
        with open(file_path, 'r') as file:
            contents = file.read()
            if error_str in contents:
                return True
            else:
                return False
    except FileNotFoundError:
        print(f"File not found: {file_path}")
        return False
    except Exception as e:
        print(f"An error occurred: {e}")
        return False

from main import P_model_list, V_model_list 

def mqtt_cmd_test(folder,device_id,device_model,sn) :

    topic_send=f"/{device_id}/send"
    topic_recv= f"/{device_id}/recv"
    mqtt_subscribe(topic_send)
    GetComponentList_cmd=read_specific_cell("./preprocess_cmd.csv", 1, 2)
    mqtt_publish(topic_recv,GetComponentList_cmd)
    time.sleep(1)
    DeviceComponentList =get_DeviceComponentList() 
    json_CompontentList = json.loads(DeviceComponentList[4:])
    Otdr_slot = OpticSwitcher_slot = Rcu_slot = PowerMeter_slot = "AD15"
    if json_CompontentList["_type"] == "DeviceComponentList" :
        componentList  = json_CompontentList.get("componentList",[])
        for component in componentList : 
            if component["_type"] == "Otdr" :
                Otdr_slot = component["slot"]
            elif component["_type"] == "OpticSwitcher" :
                OpticSwitcher_slot = component["slot"]
            elif component["_type"] == "Rcu" :
                Rcu_slot = component["slot"]
            elif component["_type"] == "PowerMeter" :
                PowerMeter_slot = component["slot"]
     
    StopReportOpticalPower_cmd=read_specific_cell("./preprocess_cmd.csv", 2, 2)
    StopReportOpticalPower_cmd = set_slot(StopReportOpticalPower_cmd,Otdr_slot,OpticSwitcher_slot,Rcu_slot,PowerMeter_slot)
    mqtt_publish(topic_recv,StopReportOpticalPower_cmd)
    
    if device_model in P_model_list:
        csv_list = ["get_cmd","P_set_cmd","task_cmd"]
    elif device_model in V_model_list :
        csv_list = ["get_cmd","V_set_cmd","task_cmd"]
    else :
        print("型号不在测试列表中！！！！！！")
    for csv_file in csv_list :
        time.sleep(0.5)
        file_path = f"./{csv_file}.csv"
        row_number = get_row_number(file_path)
        for row in range(1,row_number):
            msg_recv = read_specific_cell(file_path, row, 2)                
            msg_send_std = read_specific_cell(file_path, row, 3)
            msg_recv = set_slot(msg_recv,Otdr_slot,OpticSwitcher_slot,Rcu_slot,PowerMeter_slot)
            flag = get_flag()
            mqtt_publish(topic_recv,msg_recv)
            time.sleep(0.3)
            count = 0
            while  flag == get_flag() and count<4 :
                count = count + 1
                time.sleep(0.2)
            else:
                log_file_path = f"./test_device/{folder}/test_log.txt"  
                if count <4 :                            
                    msg_send = get_mqtt_msg()
                    json_send = json.loads(msg_send[4:])
                    json_send_std = json.loads(msg_send_std[4:])
                    if json_send["_type"]==json_send_std["_type"] :
                        print(f"{csv_file} 用例{row} ：测试成功\n")
                        try:
                            with open(log_file_path, 'a') as file:
                                file.write(f"{csv_file} 用例{row} ：测试成功\n")
                        except Exception as e:
                            print(f"An error occurred: {str(e)}")
                    else :
                        print(f"{csv_file} 用例{row} ：测试失败\n")
                        try:
                            with open(log_file_path, 'a') as file:
                                file.write(f"{csv_file} 用例{row} ：测试失败\n")
                                file.write(f"发送的命令 ：\n{msg_recv}\n")
                                file.write(f"接收到的回复 ：\n{msg_send}\n")
                        except Exception as e:
                            print(f"An error occurred: {str(e)}")

                else  :
                    print(f"ERROR: {csv_file} 用例{row} ：测试超时无回复\n")
                    try:
                        with open(log_file_path, 'a') as file:
                            file.write(f"ERROR: {csv_file} 用例{row} ：测试超时无回复\n")
                            file.write(f"发送的命令 ：\n{msg_recv}\n")
                            file.write(f"接收到的回复 ：\n{msg_send}\n")
                    except Exception as e:
                        print(f"An error occurred: {str(e)}")
            json_recv = json.loads(msg_recv[4:])
            if json_recv["_type"] == "SetTime" :
                msg_recv = get_nowtimestamp_cmd(msg_recv)
                mqtt_publish(topic_recv,msg_recv)
                time.sleep(3)                
    print(f"{folder} test end\n")
    source_file ="./report_templet.docx"
    copy_path = f"./test_device/{folder}/{folder}.docx"
    copy_file(source_file, copy_path)
    doc = Document(copy_path)
    store_time = datetime.now().date()
    test_time = store_time + timedelta(days=3)
    find_replace_in_table(doc,"store_time",f"{store_time}")
    find_replace_in_table(doc,"test_time",f"{test_time}")
    find_replace_in_table(doc,"model",device_model)
    find_replace_in_table(doc,"sn",sn)
    find_replace(doc,"test_man","聂侯辉")
    find_replace(doc,"test_time",f"{test_time}")
    testend_time = store_time + timedelta(days=9)
    find_replace(doc,"testend_time",f"{testend_time}")
    doc.save(copy_path)

def check_test_result(folder) :
    docx_path = f"./test_device/{folder}/{folder}.docx"
    doc = Document(docx_path)
    log_file_path = f"./test_device/{folder}/test_log.txt"  
    result_file_path = f"./test_device/test_result.txt"
    if check_for_ERROR(log_file_path,"ERROR") :
        result_f =  open(result_file_path, 'a') 
        result_f.write(f"ERROR：{folder}测试失败\n")
        if check_for_ERROR(log_file_path,"ERROR: task_cmd 用例2") :
            find_replace_in_table(doc,"$osw","否")
    else :
        result_f =  open(result_file_path, 'a')
        result_f.write(f"{folder}测试成功\n")
        find_replace_in_table(doc,"$osw","是")
        find_replace_in_table(doc,"$otdr","是")
    doc.save(docx_path)
        

import shutil

def copy_file(source_path, destination_path):
    try:
        shutil.copy2(source_path, destination_path)
        print(f"文件 {source_path} 已成功复制到 {destination_path}")
    except FileNotFoundError:
        print(f"文件 {source_path} 未找到。")
    except PermissionError:
        print(f"没有权限复制文件 {source_path}。")
    except Exception as e:
        print(f"复制文件时发生错误：{e}")


from docx import Document

def find_replace_in_table(doc, search_str, replace_str):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 使用 run 对象进行字符串替换
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace(search_str, replace_str)

def find_replace(doc, search_str, replace_str):
    for paragraph in doc.paragraphs:
        if search_str in paragraph.text:
            # 使用 run 对象进行字符串替换
            for run in paragraph.runs:
                run.text = run.text.replace(search_str, replace_str)